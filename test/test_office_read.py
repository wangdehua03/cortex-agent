"""Smoke test for read_word and read_pdf tools."""
import os
import sys
from pathlib import Path

# Ensure project root is on path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from src.utils.function import (
    read_docx_to_text,
    read_pdf_to_text,
    safe_path,
    _ask_path_permission,
)


def _fixture_path(filename: str) -> Path:
    return Path(__file__).resolve().parent / "fixtures" / filename


def test_docx():
    import docx
    fixture_dir = Path(__file__).resolve().parent / "fixtures"
    fixture_dir.mkdir(exist_ok=True)
    path = fixture_dir / "test_docx.docx"

    doc = docx.Document()
    doc.add_heading("Hello DOCX", level=1)
    doc.add_paragraph("This is a test paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "C"
    table.rows[1].cells[1].text = "D"
    doc.save(path)

    text = read_docx_to_text(str(path))
    assert "Hello DOCX" in text, text
    assert "This is a test paragraph." in text, text
    assert "A" in text and "B" in text, text
    print("DOCX test passed.")
    path.unlink(missing_ok=True)


def test_pdf():
    from pypdf import PdfWriter
    fixture_dir = Path(__file__).resolve().parent / "fixtures"
    fixture_dir.mkdir(exist_ok=True)
    path = fixture_dir / "test_pdf.pdf"

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)
    with open(path, "wb") as f:
        writer.write(f)

    text = read_pdf_to_text(str(path))
    assert "Page 1" in text, text
    assert "Page 2" in text, text

    text_limited = read_pdf_to_text(str(path), max_pages=1)
    assert "Page 1" in text_limited, text_limited
    assert "truncated after 1 of 2 pages" in text_limited, text_limited
    print("PDF test passed.")
    path.unlink(missing_ok=True)


def test_safe_path_permission_allowed():
    """外部路径：用户允许时应返回解析后的 Path。"""
    import src.utils.function as _fn

    original = _fn._ask_path_permission
    try:
        _fn._ask_path_permission = lambda p, r: True
        external = "C:/Windows/System32/drivers/etc/hosts" if sys.platform == "win32" else "/etc/hosts"
        result = safe_path(external)
        assert result.resolve() == Path(external).resolve(), result
        print("safe_path allow test passed.")
    finally:
        _fn._ask_path_permission = original


def test_safe_path_permission_denied():
    """外部路径：用户拒绝时应抛出 ValueError。"""
    import src.utils.function as _fn

    original = _fn._ask_path_permission
    try:
        _fn._ask_path_permission = lambda p, r: False
        external = "C:/Windows/System32/drivers/etc/hosts" if sys.platform == "win32" else "/etc/hosts"
        try:
            safe_path(external)
            raise AssertionError("expected ValueError")
        except ValueError as e:
            assert "escapes workspace" in str(e), e
        print("safe_path deny test passed.")
    finally:
        _fn._ask_path_permission = original


if __name__ == "__main__":
    test_docx()
    test_pdf()
    test_safe_path_permission_allowed()
    test_safe_path_permission_denied()
    print("All office read tests passed.")
